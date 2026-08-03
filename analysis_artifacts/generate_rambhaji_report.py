from __future__ import annotations

import csv
import math
import os
import re
from collections import Counter, defaultdict
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Pt


REPORT_DATE = "July 28, 2026"
ROOT = Path(__file__).resolve().parents[1]
FRONTEND = ROOT / "Frontend"
BACKEND = ROOT / "Backend"
OUTPUT_DIR = ROOT / "analysis_artifacts" / "output"
REPORT_PATH = OUTPUT_DIR / "Rambhaji_Project_Analysis_Hinglish.docx"
ER_SVG_PATH = OUTPUT_DIR / "Rambhaji_ER_Diagram.svg"
ER_MMD_PATH = OUTPUT_DIR / "Rambhaji_ER_Diagram.mmd"
INVENTORY_TXT_PATH = OUTPUT_DIR / "Rambhaji_Full_Inventory.txt"
INVENTORY_CSV_PATH = OUTPUT_DIR / "Rambhaji_Full_Inventory_With_Notes.csv"


@dataclass(frozen=True)
class Entity:
    name: str
    fields: list[str]
    note: str


ENTITIES = [
    Entity("User", ["id PK", "role", "wallet_balance", "due_amount", "postpaid_debt"], "Primary actor table for user, admin, and delivery roles."),
    Entity("Address", ["id PK", "user_id FK", "zone", "is_default"], "Customer address book plus delivery zone tagging."),
    Entity("Unit", ["id PK", "name", "symbol", "is_active"], "Normalizes unit definitions like kg, gm, litre, ml, piece."),
    Entity("Product", ["id PK", "unit_id FK", "category", "selling_price", "stock_qty"], "Master catalog for vegetables, fruits, water items, and retail goods."),
    Entity("Package", ["id PK", "type", "price", "services", "target_user_id"], "Monthly or yearly plans, including custom targeted plans."),
    Entity("PackageFixedItem", ["id PK", "package_id FK", "product_id FK", "qty", "unit"], "Fixed items that always go inside a package service."),
    Entity("PackageSeasonalPool", ["id PK", "package_id FK", "product_id FK"], "Allowed product pool for seasonal selection."),
    Entity("PackageSeasonalConfig", ["id PK", "package_id FK", "margin_percent", "seasonal_budget_ratio"], "Controls seasonal budget math."),
    Entity("Subscription", ["id PK", "user_id FK", "package_id FK", "address_id FK", "batch_id FK"], "Main recurring package subscription table."),
    Entity("SubscriptionItem", ["id PK", "subscription_id FK", "product_id FK", "qty", "selection_type"], "Stores fixed or seasonal product selections at subscription level."),
    Entity("WaterSubscription", ["id PK", "user_id FK", "address_id FK", "batch_id FK", "water_product_id"], "Recurring bottled water subscription."),
    Entity("PauseLog", ["id PK", "subscription_id FK", "water_subscription_id FK", "pause_start", "pause_end"], "Tracks pause windows for package and water flows."),
    Entity("Batch", ["id PK", "name", "status", "is_deleted"], "Delivery time-slot or route bucket used in assignment."),
    Entity("DeliverySchedule", ["id PK", "subscription_id FK", "water_subscription_id FK", "batch_id FK", "status"], "Daily service occurrence generated from subscriptions."),
    Entity("DeliveryItem", ["id PK", "schedule_id FK", "product_id FK", "packed_qty", "delivered_qty"], "Packed, delivered, missed, or returned items against a schedule."),
    Entity("ScheduleSeasonalSelection", ["id PK", "schedule_id FK", "product_id FK", "qty", "extra_charge"], "Future per-schedule seasonal overrides."),
    Entity("WalletTransaction", ["id PK", "user_id FK", "type", "amount", "status"], "Wallet credits and debits."),
    Entity("CreditLog", ["id PK", "user_id FK", "subscription_id FK", "short_amount", "is_overridden"], "Credit-based recovery and override records."),
    Entity("Notification", ["id PK", "user_id FK", "title", "message", "is_read"], "User-facing notification feed."),
    Entity("PaymentTransaction", ["id PK", "user_id FK", "type", "payment_method", "status"], "Payment gateway and payment-intent ledger."),
    Entity("CalculatorDraft", ["id PK", "user_id FK", "name", "type", "totals"], "Admin calculator snapshots."),
    Entity("CalculatorDraftItem", ["id PK", "draft_id FK", "product_id FK", "qty", "cost"], "Line items under calculator drafts."),
    Entity("PurchaseLog", ["id PK", "product_id FK", "qty", "purchase_price", "purchase_date"], "Stock purchase history."),
    Entity("RetailOrder", ["id PK", "user_id FK", "address_id FK", "delivery_boy_id FK", "batch_id FK"], "One-time retail checkout orders."),
    Entity("RetailOrderItem", ["id PK", "retail_order_id FK", "product_id FK", "qty", "price"], "Line items for retail orders."),
    Entity("MissedProductLog", ["id PK", "user_id FK", "product_id FK", "qty", "delivery_schedule_id"], "Carry-forward log for packed but undelivered items."),
    Entity("ReturnedProductLog", ["id PK", "user_id FK", "product_id FK", "qty", "delivery_schedule_id"], "Carry-forward log for approved returns."),
    Entity("BatchProcessingLog", ["id PK", "batch_id FK", "product_id FK", "process_type", "time_taken_minutes"], "Working/processing logs for demand preparation."),
]


RELATIONSHIPS = [
    ("User", "Address", "1", "N"),
    ("User", "Subscription", "1", "N"),
    ("User", "WaterSubscription", "1", "N"),
    ("User", "WalletTransaction", "1", "N"),
    ("User", "CreditLog", "1", "N"),
    ("User", "Notification", "1", "N"),
    ("User", "PaymentTransaction", "1", "N"),
    ("User", "CalculatorDraft", "1", "N"),
    ("User", "RetailOrder", "1", "N"),
    ("User", "MissedProductLog", "1", "N"),
    ("User", "ReturnedProductLog", "1", "N"),
    ("Unit", "Product", "1", "N"),
    ("Package", "PackageFixedItem", "1", "N"),
    ("Package", "PackageSeasonalPool", "1", "N"),
    ("Package", "PackageSeasonalConfig", "1", "1"),
    ("Package", "Subscription", "1", "N"),
    ("Product", "PackageFixedItem", "1", "N"),
    ("Product", "PackageSeasonalPool", "1", "N"),
    ("Product", "SubscriptionItem", "1", "N"),
    ("Product", "DeliveryItem", "1", "N"),
    ("Product", "ScheduleSeasonalSelection", "1", "N"),
    ("Product", "PurchaseLog", "1", "N"),
    ("Product", "RetailOrderItem", "1", "N"),
    ("Product", "MissedProductLog", "1", "N"),
    ("Product", "ReturnedProductLog", "1", "N"),
    ("Product", "BatchProcessingLog", "1", "N"),
    ("Address", "Subscription", "1", "N"),
    ("Address", "WaterSubscription", "1", "N"),
    ("Address", "RetailOrder", "1", "N"),
    ("Batch", "Subscription", "1", "N"),
    ("Batch", "WaterSubscription", "1", "N"),
    ("Batch", "DeliverySchedule", "1", "N"),
    ("Batch", "RetailOrder", "1", "N"),
    ("Batch", "BatchProcessingLog", "1", "N"),
    ("Subscription", "SubscriptionItem", "1", "N"),
    ("Subscription", "DeliverySchedule", "1", "N"),
    ("Subscription", "PauseLog", "1", "N"),
    ("WaterSubscription", "DeliverySchedule", "1", "N"),
    ("WaterSubscription", "PauseLog", "1", "N"),
    ("DeliverySchedule", "DeliveryItem", "1", "N"),
    ("DeliverySchedule", "ScheduleSeasonalSelection", "1", "N"),
    ("RetailOrder", "RetailOrderItem", "1", "N"),
    ("CalculatorDraft", "CalculatorDraftItem", "1", "N"),
    ("Product", "CalculatorDraftItem", "1", "N"),
]


FRONTEND_ROUTE_GROUPS = {
    "Public/Auth Routes": [
        "/login",
        "/register",
        "/verify-otp",
        "/forgot-password",
        "/unauthorized",
    ],
    "User Routes": [
        "/dashboard",
        "/packages",
        "/my-subscriptions",
        "/wallet",
        "/water",
        "/notifications",
        "/deliveries",
        "/addresses",
        "/retail-store",
        "/my-retail-orders",
        "/payment-status",
        "/preferences",
    ],
    "Admin Routes": [
        "/admin",
        "/admin/products",
        "/admin/units",
        "/admin/packages",
        "/admin/subscriptions",
        "/admin/users",
        "/admin/user-history",
        "/admin/user-addresses",
        "/admin/user-preferences",
        "/admin/deliveries",
        "/admin/returns",
        "/admin/demands",
        "/admin/seasonal-selections",
        "/admin/summary",
        "/admin/calculator",
        "/admin/reverse-calculator",
        "/admin/retail-orders",
        "/admin/batches",
        "/admin/all-orders",
        "/admin/missed-products",
        "/admin/product-sales",
        "/admin/working-logs",
    ],
    "Delivery Routes": [
        "/delivery",
        "/delivery/history",
    ],
}


BACKEND_ENDPOINT_GROUPS = {
    "Auth": [
        "POST /api/auth/register",
        "POST /api/auth/login",
        "GET /api/auth/me",
        "PUT /api/auth/me/dislikes",
        "POST /api/auth/logout",
        "POST /api/auth/verify-registration-otp",
        "POST /api/auth/resend-otp",
        "POST /api/auth/forgot-password",
        "POST /api/auth/verify-forgot-password-otp",
        "POST /api/auth/reset-password",
    ],
    "Addresses": [
        "PATCH /api/addresses/:id/location",
        "POST /api/addresses/admin/create",
        "GET /api/addresses/admin/unassigned",
        "PATCH /api/addresses/admin/:id/zone",
        "POST /api/addresses",
        "GET /api/addresses",
        "PUT /api/addresses/:id",
        "DELETE /api/addresses/:id",
        "PATCH /api/addresses/:id/default",
    ],
    "Products & Units": [
        "GET /api/products/public/vegetables",
        "GET /api/products",
        "POST /api/products",
        "PUT /api/products/:id",
        "PUT /api/products/:id/retail-price",
        "DELETE /api/products/:id",
        "POST /api/products/purchase",
        "GET /api/products/purchases",
        "GET /api/products/stock-summary",
        "GET /api/products/sales",
        "GET /api/units",
        "POST /api/units",
        "PUT /api/units/:id",
        "DELETE /api/units/:id",
    ],
    "Packages & Subscriptions": [
        "GET /api/packages",
        "GET /api/packages/:id",
        "POST /api/packages",
        "PUT /api/packages/:id",
        "POST /api/subscribe",
        "GET /api/my-subscriptions",
        "GET /api/available-dates",
        "POST /api/confirm-start-date",
        "PATCH /api/subscriptions/:id/pause",
        "PATCH /api/subscriptions/:id/restart",
        "PATCH /api/subscriptions/:id/cancel",
        "GET /api/seasonal-options/:subscription_id",
        "POST /api/select-seasonal",
        "PATCH /api/update-seasonal",
        "GET /api/subscriptions/:id/upcoming-selections",
        "POST /api/subscriptions/:id/schedule-seasonal",
        "POST /api/admin/assign",
        "POST /api/admin/renew",
        "PATCH /api/admin/subscriptions/:id/batch",
    ],
    "Water": [
        "POST /api/water/subscribe",
        "GET /api/water/subscriptions",
        "GET /api/water/available-dates",
        "POST /api/water/confirm-start-date",
        "PATCH /api/water/:id/pause",
        "PATCH /api/water/:id/restart",
        "PATCH /api/water/:id/cancel",
        "PATCH /api/water/admin/water/:id/batch",
    ],
    "Wallet, Notifications & Analytics": [
        "GET /api/wallet",
        "POST /api/add-funds",
        "GET /api/wallet/transactions",
        "GET /api/notifications",
        "PATCH /api/notifications/:id/mark-read",
        "GET /api/admin/user-analytics/users",
        "GET /api/admin/user-analytics/:userId",
    ],
    "Delivery & Returns": [
        "GET /api/today-deliveries",
        "POST /api/mark-delivered",
        "POST /api/return-item",
        "PATCH /api/return-item/:id/review",
        "POST /api/admin/return-item",
        "POST /api/admin/return-order",
        "GET /api/delivery-history",
        "GET /api/admin/deliveries",
        "GET /api/admin/returns",
        "GET /api/admin/demands",
        "GET /api/admin/seasonal-selections",
        "GET /api/admin/orders",
        "PUT /api/admin/orders/assign-batch",
        "PUT /api/admin/orders/assign-delivery-boy",
        "PUT /api/admin/orders/pack",
        "GET /api/admin/missed-products",
        "GET /api/available-orders",
        "PUT /api/accept-order",
        "GET /api/boy-history",
        "POST /api/boy-return-item",
        "POST /api/boy-return-order",
    ],
    "Retail": [
        "POST /api/retail",
        "GET /api/retail",
        "GET /api/retail/admin",
        "PATCH /api/retail/admin/:id/status",
    ],
    "Batches, Dashboard & Calculator": [
        "GET /api/user/batches",
        "POST /api/admin/batches",
        "GET /api/admin/batches",
        "PUT /api/admin/batches/:id",
        "DELETE /api/admin/batches/:id",
        "GET /api/admin/batches/:id/demands",
        "POST /api/admin/batches/:id/demands/process",
        "GET /api/admin/processing-logs",
        "GET /api/admin/dashboard-stats",
        "POST /api/calculator/drafts",
        "GET /api/calculator/drafts",
        "DELETE /api/calculator/drafts/:id",
    ],
    "Payment": [
        "POST /api/payment/phonepe/initiate",
        "GET /api/payment/phonepe/status/:txnId",
        "POST /api/payment/phonepe/callback",
    ],
}


FLOW_SECTIONS = [
    (
        "1. App Bootstrap Flow",
        [
            "Frontend side par `src/main.jsx` React app ko mount karta hai aur `src/App.jsx` route tree define karta hai. App start hote hi Zustand auth store ke through token-based session restore hota hai.",
            "Axios instance `src/api/axios.js` har request me `Authorization: Bearer <token>` inject karta hai. Agar backend 401 bhej de to token localStorage se hata diya jata hai aur browser `/login` par redirect hota hai.",
            "Backend side par `Backend/index.js` environment load karta hai, Express boot karta hai, CORS lagata hai, `/uploads` static serve karta hai, DB connect karta hai, Sequelize sync chalata hai, seed logic run karta hai, aur nightly cron scheduler start karta hai.",
        ],
    ),
    (
        "2. Authentication Flow",
        [
            "Register page par user name, mobile, email, password aur disliked vegetables bhej sakta hai. Backend `auth.controller.js` uniqueness check karta hai, password hash karta hai, OTP generate karta hai, aur user ko unverified state me save karta hai.",
            "Current implementation me OTP effectively demo mode me hai kyunki hardcoded `123456` use ho raha hai. Isse testing aasaan hoti hai, lekin production security weak ho jati hai.",
            "OTP verify hone ke baad login allow hota hai. Login par JWT token response aur cookie dono diye jate hain. `requireAuth` middleware token verify karta hai aur inactive users ko block karta hai.",
        ],
    ),
    (
        "3. Package Purchase + Start Date Flow",
        [
            "User `PackagesPage.jsx` me packages load karta hai, address choose karta hai, batch select karta hai, aur payment method choose karta hai.",
            "Wallet payment path me backend subscription intent validate karta hai, par delivery ke time per-service debit logic run hota hai. Online payment path me PhonePe transaction create hota hai; success par backend wallet ko credit karta hai aur actual recurring subscription create karta hai.",
            "Subscription create hone ke baad user available start dates dekh kar `confirm-start-date` hit karta hai. Tab `generateDeliveryDates` logic ke basis par future `DeliverySchedule` rows banti hain.",
        ],
    ),
    (
        "4. Seasonal Selection Flow",
        [
            "Package ke fixed items alag rehte hain aur seasonal budget alag calculate hota hai. Budget math `scheduleEngine.js` aur `subscription.controller.js` dono me enforce hota hai.",
            "User default seasonal preference subscription level par save kar sakta hai, aur future unlocked schedules ke liye alag-alag date-specific selections bhi save kar sakta hai.",
            "Nightly cron next-day schedules lock kar deta hai. Lock ke baad seasonal edits stop ho jati hain. Agar user ne kuch select nahi kiya to cron popularity-based auto-selection ya package pool fallback apply karta hai.",
        ],
    ),
    (
        "5. Water Subscription Flow",
        [
            "Water page package flow jaisa hi hai, lekin plan product-driven hota hai. Controller category `water` products me matching bottle/container variant dhoondh kar subscription banata hai.",
            "Daily water plan me 30 services aur alternate plan me 15 services assume ki jati hain. Sunday adjustment ke saath dates generate hoti hain.",
            "Pause, restart, cancel aur batch reassignment ka parallel flow water subscriptions par bhi available hai.",
        ],
    ),
    (
        "6. Retail Order Flow",
        [
            "Retail store page active products load karti hai, cart banati hai, quantity control karti hai, aur checkout par COD, wallet, ya online payment choose karne deti hai.",
            "Backend retail controller minimum retail quantity validate karta hai, quantity ko unit normalization ke hisab se convert karta hai, fixed delivery charge add karta hai, aur 8 PM cutoff ke hisab se delivery date nikalta hai.",
            "Online retail payment me pending `RetailOrder` pehle create hota hai, phir PhonePe success par payment status success ho jata hai. Admin delivery status ko delivered mark kare to stock reduction aur COD settlement hota hai.",
        ],
    ),
    (
        "7. Delivery Allocation + Packing Flow",
        [
            "Admin `AdminAllOrders.jsx` me ek date ke sabhi package, water, aur retail orders grouped form me dekhta hai. Batch assignment aur delivery boy assignment bulk mode me ho sakta hai.",
            "Packing ke waqt packed quantity product-wise di jati hai. Backend available qty ko multiple schedules aur retail items me allocate karta hai. Jo package items pack nahi hue unke liye next service carry-forward logic aur `MissedProductLog` create hota hai.",
            "Batch aur address zone ke basis par delivery boy auto-assignment bhi hota hai. Order ready hone ke baad delivery boys `available-orders` me same-user grouped work accept kar sakte hain.",
        ],
    ),
    (
        "8. Delivery Completion Flow",
        [
            "Delivery boy `today-deliveries` se schedule ya retail order open karta hai aur delivered mark karta hai. Package/water schedule par backend ensure karta hai ki corresponding `DeliveryItem` rows exist karein.",
            "Package delivery ke waqt per-service wallet debit hota hai. Agar wallet kam ho to due amount me shift hota hai. Agar actual purchase cost lower nikle to seasonal budget refund-type adjustment bhi ho sakta hai.",
            "Final loyalty-eligible service ke baad system ek extra postpaid serving create kar sakta hai aur `postpaid_debt` track karta hai. Iske baad later payment flows us debt ko recover karte hain.",
        ],
    ),
    (
        "9. Return Flow",
        [
            "User delivered history se item-level return request bhej sakta hai aur photo upload kar sakta hai. Admin review page se approve ya reject karta hai.",
            "Approve hone par stock wapas add hota hai, notification jati hai, aur `ReturnedProductLog` ke through next schedule me compensation ya carry-forward logic active hota hai.",
            "Admin aur delivery boy dono ke paas direct return endpoints bhi hain. Pura order return hone par backend ek naya future schedule create karke service recover karta hai.",
        ],
    ),
    (
        "10. Wallet, Credit, and Payment Flow",
        [
            "Wallet page balance aur transaction ledger dikhati hai. `add-funds` endpoint currently direct mock wallet credit karta hai.",
            "Package aur retail ke online payments `payment.controller.js` ke PhonePe flow se jate hain. Environment variables missing hon to controller simulated success mode me bhi kaam kar leta hai.",
            "2 paid months ke baad low-balance customers ke liye `CreditLog` create ho sakta hai. Admin override kare to paused subscription wapas activate ki ja sakti hai.",
        ],
    ),
    (
        "11. Cron and Automation Flow",
        [
            "Nightly cron 8:00 PM IST par chalne ke liye configured hai. Ye expired pauses ko auto-complete karta hai aur restart ke baad missing schedules regenerate kar sakta hai.",
            "Cron agle din ki pending unlocked schedules ke liye seasonal choices auto-fill karta hai, delivery items create karta hai, reminder notifications bhejta hai, aur un schedules ko lock kar deta hai.",
            "Wallet low-balance notifications, overdue blockage, aur credit-generation logic bhi isi scheduled automation me bundled hai.",
        ],
    ),
    (
        "12. Admin Analytics and Operations Flow",
        [
            "Admin dashboard high-level counters dikhata hai. Admin user history page ek user ke package subscriptions, water subscriptions, retail orders, delivered products, batch mapping, aur renewals ko join karke analytics surface karti hai.",
            "Admin calculator pages pricing experiments save kar sakti hain. `BatchProcessingLog` aur Working Logs pages demand preparation ka operational visibility deti hain.",
            "Package creation me fixed-cost vs per-service budget validation hota hai, isliye admin price aur composition ke beech margin-safe design maintain kar sakta hai.",
        ],
    ),
]


OBSERVATIONS = [
    "Category module half-wired lag raha hai: `category.route.js` aur `category.controller.js` legacy style me hain, lekin backend boot me mount bhi nahi hua aur imported auth helpers/pool exports current code se match nahi karte.",
    "Dashboard controller kuch counters ke liye `packing` aur `ready` statuses use kar raha hai, jabki main `DeliverySchedule` status enum `ready_for_delivery` use karta hai. Is wajah se admin dashboard counts drift kar sakte hain.",
    "Delivery controller ke retail-delivered branch me `ro.total_price` reference dikhta hai, lekin model side par `RetailOrder` amount field `total_amount` hai. Yahan settlement bug ka risk hai.",
    "PaymentTransaction enums aur controller usage fully aligned nahi dikh rahe. Code me `extra_overage_charge` aur `admin_assigned` jaise values ka flow hai jo schema enum mismatch create kar sakta hai.",
    "Water batch update route ka final URL awkward hai: `/api/water/admin/water/:id/batch`. Functionally chal sakta hai, lekin naming consistency weak hai.",
    "OTP `123456` hardcoded hone ki wajah se auth flow demo/test friendly hai, lekin real production security ke liye isko replace karna hoga.",
    "Frontend `Landing/Landing.jsx` aur `src/App.css` legacy/unused remnants lagte hain. Root routing ab `/login` par redirect karti hai, isliye Landing page active app flow ka part nahi hai.",
]


EXACT_FILE_NOTES = {
    "Frontend/.env": "Frontend environment config file. Yahan API base URL jaisi runtime values store hoti hain.",
    "Frontend/.gitignore": "Git ignore rules for frontend workspace artifacts.",
    "Frontend/.htaccess": "Apache deployment helper. HTTPS force aur SPA fallback routing set karta hai.",
    "Frontend/README.md": "Mostly default Vite README. Current business logic documentation yahan maintained nahi hai.",
    "Frontend/package.json": "Frontend dependency manifest. React 19, React Router, Zustand, Axios, Lucide, Tailwind, Vite toolchain define karta hai.",
    "Frontend/package-lock.json": "Exact npm dependency lockfile.",
    "Frontend/vite.config.js": "Vite dev/build configuration.",
    "Frontend/tailwind.config.js": "Custom color palette, fonts, and animations define karta hai.",
    "Frontend/postcss.config.js": "PostCSS/Tailwind processing configuration.",
    "Frontend/eslint.config.js": "Linting rules configuration.",
    "Frontend/theme_converter.cjs": "Theme/class conversion helper script.",
    "Frontend/theme_refiner.cjs": "Theme cleanup/refinement helper script.",
    "Frontend/index.html": "Vite entry HTML shell.",
    "Frontend/public/favicon.svg": "Static favicon asset.",
    "Frontend/public/icons.svg": "Reusable icon sprite/static icon asset.",
    "Frontend/src/main.jsx": "React bootstrap entry point.",
    "Frontend/src/App.jsx": "Complete route map aur role-based app composition ka central file.",
    "Frontend/src/App.css": "Legacy Vite starter CSS style file; app ka primary styling source nahi lagta.",
    "Frontend/src/index.css": "Shared utility classes, buttons, cards, nav styling, and base app theme yahin defined hai.",
    "Frontend/src/api/axios.js": "Central Axios client with token injection and 401 redirect handling.",
    "Frontend/src/store/authStore.js": "Zustand auth store; login/register/OTP/reset/me/logout actions yahin se manage hote hain.",
    "Frontend/src/components/Layout.jsx": "Responsive shell layout with sidebar + mobile toggle.",
    "Frontend/src/components/ProtectedRoute.jsx": "Token guard, loading state, role access control, aur public route redirection handle karta hai.",
    "Frontend/src/components/Sidebar.jsx": "Role-based navigation menu definitions aur logout button render karta hai.",
    "Frontend/src/pages/LoginPage.jsx": "User/admin/delivery login UI with role-specific redirect.",
    "Frontend/src/pages/RegisterPage.jsx": "Registration form + disliked vegetables preference capture.",
    "Frontend/src/pages/VerifyOTPPage.jsx": "Registration OTP verification + resend timer flow.",
    "Frontend/src/pages/ForgotPasswordPage.jsx": "Forgot password, OTP verify, reset password multi-step screen.",
    "Frontend/src/pages/Landing/Landing.jsx": "Legacy landing page file. Current main router se primary flow me active nahi dikh raha.",
    "Frontend/src/pages/user/PackagesPage.jsx": "User ka main package browsing, address select, subscribe, PhonePe initiate, start-date, and seasonal selection page.",
    "Frontend/src/pages/user/MySubscriptions.jsx": "Existing subscriptions ka management page: pause, restart, cancel, start date confirm, default seasonal, date-wise seasonal edits.",
    "Frontend/src/pages/user/WalletPage.jsx": "Wallet balance, add funds, transaction history UI.",
    "Frontend/src/pages/user/WaterPage.jsx": "Water subscription purchase and lifecycle management UI.",
    "Frontend/src/pages/user/NotificationsPage.jsx": "Notification list and mark-as-read page.",
    "Frontend/src/pages/user/DeliveryHistory.jsx": "Delivered order history and return request flow.",
    "Frontend/src/pages/user/AddressPage.jsx": "Address CRUD, default address, and user-side zone-neutral address management.",
    "Frontend/src/pages/user/RetailStore.jsx": "One-time retail shopping cart and checkout flow.",
    "Frontend/src/pages/user/MyRetailOrders.jsx": "Retail order history/status page.",
    "Frontend/src/pages/user/PaymentStatusPage.jsx": "PhonePe payment status polling and finalization screen.",
    "Frontend/src/pages/user/UserDashboard.jsx": "User overview dashboard with quick links and summary widgets.",
    "Frontend/src/pages/user/UserPreferences.jsx": "Disliked vegetables preference editor.",
    "Frontend/src/pages/admin/AdminDashboard.jsx": "Admin summary counters and user overview data screen.",
    "Frontend/src/pages/admin/AdminProducts.jsx": "Product CRUD, purchases, stock summary, and retail price management UI.",
    "Frontend/src/pages/admin/AdminUnits.jsx": "Measurement unit master maintenance page.",
    "Frontend/src/pages/admin/AdminPackages.jsx": "Package builder/editor with calculator draft support.",
    "Frontend/src/pages/admin/AdminSubscriptions.jsx": "Admin-side user subscription listing page.",
    "Frontend/src/pages/admin/AdminUsers.jsx": "Admin user list, status toggle, and user creation page.",
    "Frontend/src/pages/admin/AdminUserHistory.jsx": "Single user ka deep subscription, delivery, and ordering analytics page.",
    "Frontend/src/pages/admin/AdminUserAddresses.jsx": "Admin-created addresses for users and address management helper screen.",
    "Frontend/src/pages/admin/AdminUserPreferences.jsx": "Admin view se disliked vegetables/preferences manage karne ka screen.",
    "Frontend/src/pages/admin/AdminDeliveries.jsx": "Delivered orders overview and admin return actions.",
    "Frontend/src/pages/admin/AdminReturns.jsx": "Pending/processed returns review UI.",
    "Frontend/src/pages/admin/AdminDemands.jsx": "Date-wise and batch-wise demand aggregation page.",
    "Frontend/src/pages/admin/AdminSeasonalSelections.jsx": "Upcoming schedules ke seasonal picks aur fallback status dekhne ka admin page.",
    "Frontend/src/pages/admin/AdminRetailOrders.jsx": "Retail order operations and status updates.",
    "Frontend/src/pages/admin/AdminBatches.jsx": "Batch CRUD management page.",
    "Frontend/src/pages/admin/AdminAllOrders.jsx": "Date-wise complete order board with batch assignment, packing, and delivery assignment.",
    "Frontend/src/pages/admin/AdminMissedProducts.jsx": "Carry-forward missed products visibility page.",
    "Frontend/src/pages/admin/AdminProductSales.jsx": "Product sales analytics page.",
    "Frontend/src/pages/admin/AdminCalculator.jsx": "Forward pricing calculator page for package planning.",
    "Frontend/src/pages/admin/AdminReverseCalculator.jsx": "Reverse calculator page for target-price back solving.",
    "Frontend/src/pages/admin/WorkingLogs.jsx": "Batch processing/working logs viewer.",
    "Frontend/src/pages/delivery/DeliveryHome.jsx": "Delivery boy dashboard with today deliveries, available orders, accept flow, geotag, and mark-delivered actions.",
    "Frontend/src/pages/delivery/DeliveryBoyHistory.jsx": "Delivery boy delivered history plus return shortcuts.",
    "Frontend/src/assets/hero.png": "Landing/marketing style hero image asset.",
    "Frontend/src/assets/react.svg": "Default React asset leftover.",
    "Frontend/src/assets/vite.svg": "Default Vite asset leftover.",
    "Backend/.env": "Backend runtime configuration: DB credentials, JWT secret, frontend URL, ImageKit keys, and app port values.",
    "Backend/package.json": "Backend dependencies manifest: Express 5, Sequelize, MySQL, JWT, bcrypt, multer, cron, xlsx, ImageKit.",
    "Backend/package-lock.json": "Exact backend dependency lockfile.",
    "Backend/index.js": "Backend bootstrap file. Routes mount karta hai, DB connect/sync/seed start karta hai, and cron scheduler boot karta hai.",
    "Backend/API_Integration_Document.docx": "Existing project-side API integration document artifact.",
    "Backend/confiq/db.js": "Sequelize MySQL connection and DB connect helper.",
    "Backend/confiq/imagekit.js": "ImageKit SDK configuration.",
    "Backend/middlewares/auth.middleware.js": "JWT auth and role guard middleware.",
    "Backend/middlewares/multer.js": "Disk-based upload middleware for images, docs, and media.",
    "Backend/middlewares/upload.middleware.js": "Memory-based larger upload middleware; current routes me limited use dikh raha hai.",
    "Backend/models/index.js": "Main Sequelize schema and all associations ka single source of truth.",
    "Backend/models/model.js": "Legacy raw SQL style model stub; current runtime ORM source nahi lagta.",
    "Backend/controllers/auth.controller.js": "Registration, login, OTP verification, password reset, and preference update logic.",
    "Backend/controllers/address.controller.js": "Address CRUD, default switching, admin create, zone update, and geolocation patch logic.",
    "Backend/controllers/product.controller.js": "Product CRUD, purchases, stock summary, product sales analytics, and retail price update logic.",
    "Backend/controllers/package.controller.js": "Package CRUD plus fixed-cost-vs-service validation and custom targeting logic.",
    "Backend/controllers/subscription.controller.js": "Recurring package subscription lifecycle, schedule creation, pause/restart/cancel, seasonal selection, renewals, and batch update logic.",
    "Backend/controllers/water.controller.js": "Water subscription lifecycle logic mirroring subscription flow with water-specific rules.",
    "Backend/controllers/wallet.controller.js": "Wallet balance, add funds, tomorrow summary, admin user ops, credit override, and preferences update logic.",
    "Backend/controllers/delivery.controller.js": "Today deliveries, marking delivered, packing allocation, returns, demand aggregation, carry-forward, and order assignment ka sabse heavy operational logic.",
    "Backend/controllers/retail.controller.js": "Retail checkout, retail order history, and admin retail status logic.",
    "Backend/controllers/payment.controller.js": "PhonePe initiation, status polling, callback handling, and online payment finalization logic.",
    "Backend/controllers/dashboard.controller.js": "Admin dashboard counter aggregation logic.",
    "Backend/controllers/notification.controller.js": "Notification list and mark-as-read logic.",
    "Backend/controllers/calculator.controller.js": "Calculator draft save/list/delete handlers.",
    "Backend/controllers/batch.controller.js": "Batch CRUD, demand processing, and processing log retrieval logic.",
    "Backend/controllers/unit.controller.js": "Unit master CRUD logic.",
    "Backend/controllers/userAnalytics.controller.js": "Per-user deep analytics join logic.",
    "Backend/controllers/category.controller.js": "Legacy category controller; current active route graph se disconnected lagta hai.",
    "Backend/routes/auth.route.js": "Auth endpoints register karta hai.",
    "Backend/routes/address.route.js": "Address endpoints register karta hai.",
    "Backend/routes/product.route.js": "Product endpoints register karta hai.",
    "Backend/routes/package.route.js": "Package endpoints register karta hai.",
    "Backend/routes/subscription.route.js": "Subscription endpoints register karta hai.",
    "Backend/routes/water.route.js": "Water subscription endpoints register karta hai.",
    "Backend/routes/wallet.route.js": "Wallet/admin-user endpoints register karta hai.",
    "Backend/routes/delivery.route.js": "Delivery, packing, admin orders, returns, and delivery-boy endpoints register karta hai.",
    "Backend/routes/retail.route.js": "Retail endpoints register karta hai.",
    "Backend/routes/payment.route.js": "PhonePe payment endpoints register karta hai.",
    "Backend/routes/dashboard.route.js": "Admin dashboard route register karta hai.",
    "Backend/routes/notification.route.js": "Notification routes register karta hai.",
    "Backend/routes/calculator.route.js": "Calculator draft routes register karta hai.",
    "Backend/routes/batch.route.js": "Batch management and processing log routes register karta hai.",
    "Backend/routes/unit.routes.js": "Unit routes register karta hai.",
    "Backend/routes/userAnalytics.route.js": "User analytics routes register karta hai.",
    "Backend/routes/category.route.js": "Legacy category routes file; current exports/middleware names stale lagte hain.",
    "Backend/utils/scheduleEngine.js": "Delivery date generation aur seasonal/yearly amount math utilities.",
    "Backend/utils/cronJobs.js": "Nightly automation, reminders, auto-selection, wallet checks, and pause completion scheduler.",
    "Backend/utils/seed.js": "Initial user/product/package seed logic.",
    "Backend/utils/importSubscribers.js": "Excel-based subscriber import tool.",
    "Backend/utils/imagekitUpload.js": "ImageKit upload helper.",
    "Backend/utils/helper.js": "Legacy helper functions like OTP/token utilities.",
    "Backend/utils/runSync.js": "Targeted Sequelize sync helper for schedule seasonal selection.",
    "Backend/utils/addColumns.js": "Schema patch helper for delivery columns.",
    "Backend/add_5_time_columns.js": "Products table me process-time columns add karne wala migration script.",
    "Backend/add_column.js": "Packages table me target mobile field add karne wala migration script.",
    "Backend/add_disliked_columns.js": "Users table me disliked products field add karne wala migration script.",
    "Backend/add_loyalty_fields.js": "Loyalty/postpaid related fields add karne wala migration script.",
    "Backend/add_processing_log_table.js": "BatchProcessingLog table sync helper.",
    "Backend/add_returned_by.js": "Delivery items me returned_by field add karne wala migration script.",
    "Backend/add_time_column.js": "Legacy preparation_time field migration helper.",
    "Backend/add_zone_columns.js": "Zone/delivery assignment related schema patch script.",
    "Backend/db-update.js": "Legacy DB patch script for returned-product related columns.",
    "Backend/migrate_units.js": "Existing products ko normalized units schema me move karne ka migration helper.",
    "Backend/read_excel.js": "Excel reading helper/test script.",
    "Backend/update_processing_logs.js": "Processing log schema update script.",
}


LAYOUT = {
    "User": (40, 40),
    "Address": (40, 250),
    "WalletTransaction": (40, 460),
    "CreditLog": (40, 650),
    "Notification": (40, 840),
    "PaymentTransaction": (40, 1030),
    "CalculatorDraft": (40, 1230),
    "RetailOrder": (40, 1420),
    "Package": (430, 40),
    "PackageFixedItem": (430, 240),
    "PackageSeasonalPool": (430, 430),
    "PackageSeasonalConfig": (430, 620),
    "Subscription": (430, 810),
    "SubscriptionItem": (430, 1030),
    "PauseLog": (430, 1210),
    "WaterSubscription": (430, 1400),
    "Unit": (820, 40),
    "Product": (820, 220),
    "PurchaseLog": (820, 430),
    "CalculatorDraftItem": (820, 620),
    "DeliverySchedule": (820, 820),
    "DeliveryItem": (820, 1030),
    "ScheduleSeasonalSelection": (820, 1220),
    "Batch": (820, 1420),
    "RetailOrderItem": (1210, 40),
    "MissedProductLog": (1210, 260),
    "ReturnedProductLog": (1210, 470),
    "BatchProcessingLog": (1210, 680),
}


def human_size(size: int) -> str:
    if size < 1024:
        return f"{size} B"
    units = ["KB", "MB", "GB"]
    value = float(size)
    for unit in units:
        value /= 1024.0
        if value < 1024.0:
            return f"{value:.1f} {unit}"
    return f"{value:.1f} TB"


def chunks(items: list[str], size: int) -> Iterable[list[str]]:
    for index in range(0, len(items), size):
        yield items[index : index + size]


def scan_tree(base: Path) -> tuple[list[str], list[tuple[str, int]]]:
    directories = [base.relative_to(ROOT).as_posix()]
    files: list[tuple[str, int]] = []
    for current_root, dirnames, filenames in os.walk(base):
        dirnames.sort()
        filenames.sort()
        current_path = Path(current_root)
        for dirname in dirnames:
            rel_dir = (current_path / dirname).relative_to(ROOT).as_posix()
            directories.append(rel_dir)
        for filename in filenames:
            file_path = current_path / filename
            rel_file = file_path.relative_to(ROOT).as_posix()
            files.append((rel_file, file_path.stat().st_size))
    directories.sort()
    files.sort(key=lambda item: item[0])
    return directories, files


def env_keys(path: Path) -> list[str]:
    keys: list[str] = []
    if not path.exists():
        return keys
    for raw_line in path.read_text(encoding="utf-8", errors="ignore").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key = line.split("=", 1)[0].strip()
        if key:
            keys.append(key)
    return sorted(set(keys))


def slug_to_words(text: str) -> str:
    text = re.sub(r"([a-z0-9])([A-Z])", r"\1 \2", text)
    text = text.replace("_", " ").replace("-", " ")
    return re.sub(r"\s+", " ", text).strip().title()


def describe_directory(path_str: str) -> str:
    path = Path(path_str)
    parts = path.parts
    if "node_modules" in parts:
        return "Third-party dependency directory installed by npm."
    if "uploads" in parts:
        return "Uploaded media/documents directory used by runtime or sample data."
    if path_str == "Frontend":
        return "Complete frontend application root."
    if path_str == "Backend":
        return "Complete backend application root."
    if path.name == "src":
        return "Main source code directory."
    if path.name == "pages":
        return "Page-level UI components."
    if path.name == "components":
        return "Shared UI and layout components."
    if path.name == "api":
        return "Frontend API abstraction helpers."
    if path.name == "store":
        return "Frontend state store definitions."
    if path.name == "controllers":
        return "Express business logic handlers."
    if path.name == "routes":
        return "Express route registration files."
    if path.name == "models":
        return "Sequelize schema definitions."
    if path.name == "middlewares":
        return "Cross-cutting Express middleware."
    if path.name == "utils":
        return "Helper scripts, cron logic, and migration helpers."
    if path.name == "confiq":
        return "Runtime configuration files."
    if len(parts) >= 3 and parts[0] == "Frontend" and parts[1] == "src" and parts[2] == "pages":
        if len(parts) == 4:
            return f"{parts[3].title()} pages folder."
    return "Project directory."


def describe_file(path_str: str) -> str:
    if path_str in EXACT_FILE_NOTES:
        return EXACT_FILE_NOTES[path_str]

    path = Path(path_str)
    parts = path.parts
    suffix = path.suffix.lower()

    if "node_modules" in parts:
        if suffix in {".js", ".mjs", ".cjs", ".json", ".d.ts", ".ts"}:
            return "Third-party package runtime/build/source file installed via npm."
        return "Third-party dependency asset or metadata file."

    if "uploads" in parts:
        return "Uploaded media proof/product asset stored in backend uploads."

    if parts[0] == "Frontend":
        if path.name == ".env":
            return "Frontend environment configuration file."
        if suffix in {".svg", ".png", ".jpg", ".jpeg", ".webp"}:
            return "Static frontend asset file."
        if "pages" in parts:
            role = "general"
            if "admin" in parts:
                role = "admin"
            elif "user" in parts:
                role = "user"
            elif "delivery" in parts:
                role = "delivery"
            stem_words = slug_to_words(path.stem)
            return f"{role.title()} page component for the {stem_words} flow."
        if "components" in parts:
            return f"Shared frontend component: {slug_to_words(path.stem)}."
        if "store" in parts:
            return f"Frontend state management file: {slug_to_words(path.stem)}."
        if "api" in parts:
            return f"Frontend API helper: {slug_to_words(path.stem)}."
        if suffix in {".js", ".jsx", ".cjs"}:
            return f"Frontend source/config script: {slug_to_words(path.stem)}."
        if suffix in {".json", ".css", ".html"}:
            return "Frontend config or presentation file."

    if parts[0] == "Backend":
        if path.name == ".env":
            return "Backend environment configuration file."
        if suffix in {".png", ".jpg", ".jpeg", ".webp", ".pdf", ".doc", ".docx", ".xls", ".xlsx"}:
            return "Backend asset, uploaded file, or supporting document."
        if "controllers" in parts:
            return f"Backend controller handling {slug_to_words(path.stem.replace('.controller', ''))} logic."
        if "routes" in parts:
            return f"Express route registration file for {slug_to_words(path.stem.replace('.route', '').replace('.routes', ''))}."
        if "middlewares" in parts:
            return f"Backend middleware file for {slug_to_words(path.stem)}."
        if "models" in parts:
            return f"Database model/schema file: {slug_to_words(path.stem)}."
        if "utils" in parts:
            return f"Utility or maintenance script: {slug_to_words(path.stem)}."
        if "confiq" in parts:
            return f"Runtime configuration file: {slug_to_words(path.stem)}."
        if suffix in {".js", ".cjs"}:
            return f"Backend script/config file: {slug_to_words(path.stem)}."
        if suffix in {".json"}:
            return "Backend configuration or manifest file."

    return "Project file."


def style_document(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)

    if "CodeBlock" not in document.styles:
        code_style = document.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Consolas"
        code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        code_style.font.size = Pt(8)

    if "SmallNote" not in document.styles:
        small_style = document.styles.add_style("SmallNote", WD_STYLE_TYPE.PARAGRAPH)
        small_style.font.name = "Calibri"
        small_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        small_style.font.size = Pt(9)


def add_bullets(document: Document, items: Iterable[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_code_lines(document: Document, lines: list[str], chunk_size: int = 250) -> None:
    for group in chunks(lines, chunk_size):
        document.add_paragraph("\n".join(group), style="CodeBlock")


def authored_files_only(paths: list[tuple[str, int]]) -> list[tuple[str, int]]:
    result = []
    for path, size in paths:
        if "/node_modules/" in f"/{path}/":
            continue
        result.append((path, size))
    return result


def make_inventory_lines(directories: list[str], files: list[tuple[str, int]]) -> list[str]:
    lines = []
    for directory in directories:
        lines.append(f"[DIR ] {directory}")
    for path, size in files:
        lines.append(f"[FILE] {path} ({human_size(size)})")
    return lines


def subtree_stats(base_name: str, directories: list[str], files: list[tuple[str, int]]) -> list[str]:
    dir_counter: Counter[str] = Counter()
    file_counter: Counter[str] = Counter()
    size_counter: Counter[str] = Counter()

    base_prefix = f"{base_name}/"
    for directory in directories:
        if directory == base_name:
            continue
        rel = directory[len(base_prefix) :] if directory.startswith(base_prefix) else directory
        head = rel.split("/", 1)[0]
        dir_counter[head] += 1

    for path, size in files:
        rel = path[len(base_prefix) :] if path.startswith(base_prefix) else path
        head = rel.split("/", 1)[0]
        file_counter[head] += 1
        size_counter[head] += size

    lines = []
    for key in sorted(set(dir_counter) | set(file_counter)):
        lines.append(
            f"{base_name}/{key}: {dir_counter[key]} subdirs, {file_counter[key]} files, {human_size(size_counter[key])}"
        )
    return lines


def relationship_lines() -> list[str]:
    return [f"{left} ({lcard}) -> ({rcard}) {right}" for left, right, lcard, rcard in RELATIONSHIPS]


def build_mermaid() -> str:
    lines = ["erDiagram"]
    for left, right, lcard, rcard in RELATIONSHIPS:
        if lcard == "1" and rcard == "N":
            card = "||--o{"
        elif lcard == "1" and rcard == "1":
            card = "||--||"
        else:
            card = "}o--o{"
        lines.append(f"    {left} {card} {right} : relates")
    return "\n".join(lines) + "\n"


def draw_svg() -> str:
    box_width = 300
    line_height = 18
    padding = 16
    entity_map = {entity.name: entity for entity in ENTITIES}
    heights = {entity.name: 34 + padding + (len(entity.fields) * line_height) + 22 for entity in ENTITIES}
    width = 1560
    height = 1680

    def center_of(name: str) -> tuple[int, int]:
        x, y = LAYOUT[name]
        return x + box_width // 2, y + heights[name] // 2

    pieces = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">',
        "<defs>",
        '<style><![CDATA['
        "text { font-family: Arial, sans-serif; fill: #162129; }"
        ".title { font-size: 28px; font-weight: 700; }"
        ".subtitle { font-size: 14px; fill: #49616d; }"
        ".entity { fill: #f8fbfc; stroke: #7ea8b7; stroke-width: 2; rx: 12; ry: 12; }"
        ".header { fill: #dff1f5; }"
        ".ename { font-size: 16px; font-weight: 700; }"
        ".field { font-size: 12px; }"
        ".line { stroke: #7d8d95; stroke-width: 2; fill: none; opacity: 0.85; }"
        ".line-dim { font-size: 11px; fill: #607885; }"
        ']]></style>',
        "</defs>",
        '<rect width="100%" height="100%" fill="#eef6f8"/>',
        f'<text x="40" y="34" class="title">Rambhaji ER Diagram</text>',
        f'<text x="40" y="58" class="subtitle">Generated from code analysis on {REPORT_DATE}. PK/FK fields are intentionally trimmed to keep the diagram readable.</text>',
    ]

    for left, right, lcard, rcard in RELATIONSHIPS:
        x1, y1 = center_of(left)
        x2, y2 = center_of(right)
        mid_x = math.floor((x1 + x2) / 2)
        pieces.append(f'<path class="line" d="M{x1} {y1} L{mid_x} {y1} L{mid_x} {y2} L{x2} {y2}" />')
        label_x = mid_x + 4
        label_y = math.floor((y1 + y2) / 2) - 4
        pieces.append(f'<text x="{label_x}" y="{label_y}" class="line-dim">{escape(lcard + ":" + rcard)}</text>')

    for entity in ENTITIES:
        x, y = LAYOUT[entity.name]
        h = heights[entity.name]
        pieces.append(f'<rect class="entity" x="{x}" y="{y}" width="{box_width}" height="{h}" />')
        pieces.append(f'<rect class="entity header" x="{x}" y="{y}" width="{box_width}" height="34" />')
        pieces.append(f'<text x="{x + 14}" y="{y + 22}" class="ename">{escape(entity.name)}</text>')
        current_y = y + 54
        for field in entity.fields:
            pieces.append(f'<text x="{x + 16}" y="{current_y}" class="field">{escape(field)}</text>')
            current_y += line_height
        pieces.append(f'<text x="{x + 16}" y="{h + y - 12}" class="line-dim">{escape(entity.note)}</text>')

    pieces.append("</svg>")
    return "\n".join(pieces)


def report_intro(front_dirs: list[str], front_files: list[tuple[str, int]], back_dirs: list[str], back_files: list[tuple[str, int]]) -> list[str]:
    total_dirs = len(front_dirs) + len(back_dirs)
    total_files = len(front_files) + len(back_files)
    total_size = sum(size for _, size in front_files) + sum(size for _, size in back_files)
    return [
        f"Ye report `{REPORT_DATE}` ko `Frontend/` aur `Backend/` dono folders ke recursive analysis par based hai. Is analysis pass me actual filesystem inventory walk ki gayi hai, isliye appendix aur companion inventory files me har discovered path capture kiya gaya hai.",
        f"Combined scan result: {total_dirs} directories aur {total_files} files. Approximate combined size {human_size(total_size)} nikla.",
        "Business logic ke hisab se ye project ek 3-role fresh-delivery platform lagta hai jisme `user`, `admin`, aur `delivery` actors hain. Core domains hain: packages, seasonal selection, water subscription, retail orders, wallet/due management, delivery operations, returns, batch processing, notifications, aur analytics.",
    ]


def generate_report() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    front_dirs, front_files = scan_tree(FRONTEND)
    back_dirs, back_files = scan_tree(BACKEND)
    all_dirs = front_dirs + back_dirs
    all_files = front_files + back_files
    all_inventory_lines = make_inventory_lines(all_dirs, all_files)

    with INVENTORY_TXT_PATH.open("w", encoding="utf-8") as handle:
        handle.write("\n".join(all_inventory_lines) + "\n")

    with INVENTORY_CSV_PATH.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.writer(handle)
        writer.writerow(["type", "path", "size_bytes", "note"])
        for directory in all_dirs:
            writer.writerow(["dir", directory, "", describe_directory(directory)])
        for path, size in all_files:
            writer.writerow(["file", path, size, describe_file(path)])

    document = Document()
    style_document(document)
    document.core_properties.title = "Rambhaji Project Analysis Hinglish"
    document.core_properties.subject = "Frontend + Backend full code analysis with ER overview"
    document.core_properties.author = "OpenAI Codex"

    document.add_heading("Rambhaji Project Full Analysis (Hinglish)", level=0)
    document.add_paragraph(
        f"Prepared on {REPORT_DATE}. Source scope: complete recursive scan of `Frontend/` and `Backend/`, including business code, configs, migration scripts, uploaded assets, and installed dependency trees."
    )

    document.add_heading("1. Executive Summary", level=1)
    for paragraph in report_intro(front_dirs, front_files, back_dirs, back_files):
        document.add_paragraph(paragraph)

    document.add_heading("2. Inventory Coverage Confirmation", level=1)
    add_bullets(
        document,
        [
            f"Frontend: {len(front_dirs)} directories, {len(front_files)} files.",
            f"Backend: {len(back_dirs)} directories, {len(back_files)} files.",
            f"Exhaustive inventory exported separately to `{INVENTORY_TXT_PATH.name}` and `{INVENTORY_CSV_PATH.name}`.",
            "Word report ke appendix me bhi full inventory blocks include kiye gaye hain taaki recursive scan traceable rahe.",
        ],
    )

    document.add_heading("3. Frontend Architecture", level=1)
    frontend_paragraphs = [
        "Frontend Vite + React 19 app hai. Routing `src/App.jsx` me hai, session `src/store/authStore.js` me Zustand se manage hota hai, aur API communication central Axios client ke through hota hai.",
        "Role-wise route segregation clearly dikh raha hai: public auth pages, user journey pages, admin operations pages, aur delivery boy workflow pages. `ProtectedRoute.jsx` token aur role dono validate karta hai.",
        "UI shell `Layout.jsx` aur `Sidebar.jsx` ke through common navigation provide karta hai. Styling ka main source `src/index.css` aur Tailwind config hai; `src/App.css` mostly legacy leftover lagta hai.",
    ]
    for paragraph in frontend_paragraphs:
        document.add_paragraph(paragraph)

    document.add_heading("3.1 Frontend Route Map", level=2)
    for group_name, routes in FRONTEND_ROUTE_GROUPS.items():
        document.add_paragraph(group_name, style="List Bullet")
        add_bullets(document, routes)

    document.add_heading("4. Backend Architecture", level=1)
    backend_paragraphs = [
        "Backend Express 5 + Sequelize + MySQL stack use karta hai. `Backend/index.js` app bootstrap, route mount, static uploads serving, DB sync, seeding, aur cron startup ka center point hai.",
        "Business logic strongly controller-centric hai. Routes thin layer hain, controllers me major domain rules likhe gaye hain, aur `models/index.js` me schema + associations ka dense centralized definition hai.",
        "Operational complexity sabse zyada subscription, delivery, packing, returns, and cron automation modules me hai. Ye project sirf CRUD app nahi hai; isme inventory, service-cycle, wallet, and recovery rules ka real workflow engine type behavior hai.",
    ]
    for paragraph in backend_paragraphs:
        document.add_paragraph(paragraph)

    document.add_heading("4.1 Backend Endpoint Catalog", level=2)
    for group_name, endpoints in BACKEND_ENDPOINT_GROUPS.items():
        document.add_paragraph(group_name, style="List Bullet")
        add_bullets(document, endpoints)

    document.add_heading("5. Core Data Model", level=1)
    document.add_paragraph(
        "Database design ka center `User`, `Product`, `Package`, `Subscription`, `DeliverySchedule`, `DeliveryItem`, `RetailOrder`, aur `Batch` domain tables hain. Inke around wallet, payment, notifications, pause, carry-forward, and analytics support tables layered hain."
    )
    document.add_heading("5.1 Entity Summary", level=2)
    for entity in ENTITIES:
        document.add_paragraph(
            f"{entity.name}: {entity.note} Key fields -> {', '.join(entity.fields)}",
            style="List Bullet",
        )

    document.add_heading("5.2 Relationship Summary", level=2)
    add_bullets(document, relationship_lines())

    document.add_heading("6. End-to-End Flows", level=1)
    for title, paragraphs in FLOW_SECTIONS:
        document.add_heading(title, level=2)
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)

    document.add_heading("7. Authored File and Folder Commentary", level=1)
    document.add_paragraph(
        "Is section me business-relevant authored files ka purpose diya gaya hai. Vendor dependencies (`node_modules`) ka exhaustive path-level catalog appendix aur CSV me diya gaya hai, lekin unke internal source ko business-flow explanation me line-by-line break nahi kiya gaya kyunki wo third-party libraries hain."
    )

    authored = authored_files_only(all_files)
    frontend_authored = [(path, size) for path, size in authored if path.startswith("Frontend/")]
    backend_authored = [(path, size) for path, size in authored if path.startswith("Backend/")]

    document.add_heading("7.1 Frontend Authored Files", level=2)
    for path, size in frontend_authored:
        document.add_paragraph(f"{path} ({human_size(size)}): {describe_file(path)}", style="List Bullet")

    document.add_heading("7.2 Backend Authored Files", level=2)
    for path, size in backend_authored:
        document.add_paragraph(f"{path} ({human_size(size)}): {describe_file(path)}", style="List Bullet")

    document.add_heading("8. Folder-Level Statistics", level=1)
    document.add_paragraph("Immediate subtree roll-up niche diya gaya hai taaki quickly samajh aaye ki kaunse sections sabse heavy hain.")
    document.add_heading("8.1 Frontend Subtrees", level=2)
    add_bullets(document, subtree_stats("Frontend", front_dirs, front_files))
    document.add_heading("8.2 Backend Subtrees", level=2)
    add_bullets(document, subtree_stats("Backend", back_dirs, back_files))

    document.add_heading("9. Environment and Runtime Keys", level=1)
    document.add_paragraph("Sensitive values intentionally expose nahi kiye gaye. Sirf discovered key names list kiye gaye hain.")
    env_lines = [
        f"Frontend/.env keys: {', '.join(env_keys(FRONTEND / '.env')) or 'none found'}",
        f"Backend/.env keys: {', '.join(env_keys(BACKEND / '.env')) or 'none found'}",
    ]
    add_bullets(document, env_lines)

    document.add_heading("10. Important Observations and Risks", level=1)
    add_bullets(document, OBSERVATIONS)

    document.add_heading("11. Generated Artifacts", level=1)
    add_bullets(
        document,
        [
            f"Primary Word report: {REPORT_PATH.name}",
            f"ER diagram SVG: {ER_SVG_PATH.name}",
            f"ER diagram Mermaid source: {ER_MMD_PATH.name}",
            f"Full path inventory text: {INVENTORY_TXT_PATH.name}",
            f"Full path inventory CSV with notes: {INVENTORY_CSV_PATH.name}",
        ],
    )

    document.add_heading("Appendix A. Full Inventory (All Directories + Files)", level=1)
    document.add_paragraph(
        "Neeche recursive filesystem inventory blocks diye gaye hain. `[DIR ]` aur `[FILE]` prefixes path type indicate karte hain. Ye list scan time par discovered har path ko include karti hai."
    )
    add_code_lines(document, all_inventory_lines, chunk_size=220)

    document.add_heading("Appendix B. Relationship Catalog", level=1)
    add_code_lines(document, relationship_lines(), chunk_size=80)

    document.save(REPORT_PATH)

    ER_MMD_PATH.write_text(build_mermaid(), encoding="utf-8")
    ER_SVG_PATH.write_text(draw_svg(), encoding="utf-8")


if __name__ == "__main__":
    generate_report()
    print(f"Generated: {REPORT_PATH}")
    print(f"Generated: {ER_SVG_PATH}")
    print(f"Generated: {ER_MMD_PATH}")
    print(f"Generated: {INVENTORY_TXT_PATH}")
    print(f"Generated: {INVENTORY_CSV_PATH}")
